# src/app/services/export_service.py
# Purpose: render structured blocks to DOCX (python-docx) or PDF (reportlab),
#          preserving heading levels, list bullets, and blockquotes.
# Exports: export_document, MEDIA_TYPES

import io
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import ListFlowable, Paragraph, SimpleDocTemplate, Spacer

from app.core.blocks import Block
from app.core.exceptions import ValidationError

MEDIA_TYPES = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}

_VERA = Path(__file__).resolve().parents[3] / "reportlab" / "fonts" / "Vera.ttf"


def _register_fonts() -> str:
    if _VERA.exists():
        pdfmetrics.registerFont(TTFont("Vera", str(_VERA)))
        return "Vera"
    return "Helvetica"


def _to_docx(blocks: list[Block]) -> bytes:
    doc = DocxDocument()
    for block in blocks:
        if block.type == "heading":
            level = min(block.level or 1, 9)
            doc.add_heading(block.text, level=level)
        elif block.type == "list_item":
            doc.add_paragraph(block.text, style="List Bullet")
        elif block.type == "blockquote":
            paragraph = doc.add_paragraph(block.text)
            if paragraph.runs:
                paragraph.runs[0].italic = True
        else:
            doc.add_paragraph(block.text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _to_pdf(blocks: list[Block]) -> bytes:
    font = _register_fonts()
    body = ParagraphStyle("body", fontName=font, fontSize=10.5, leading=15, alignment=TA_LEFT)
    heading_styles = [
        ParagraphStyle(
            f"h{i}",
            parent=body,
            fontName=font,
            fontSize=18 - i * 1.5,
            leading=22 - i * 2,
            spaceAfter=8,
        )
        for i in range(1, 7)
    ]
    quote = ParagraphStyle(
        "quote", parent=body, fontName=font, leftIndent=1 * cm, textColor="#555555", spaceAfter=8
    )
    buffer = io.BytesIO()
    story: list = []
    for block in blocks:
        if block.type == "heading":
            story.append(
                Paragraph(
                    block.text.replace("&", "&amp;"),
                    heading_styles[min(block.level or 1, 6) - 1],
                )
            )
        elif block.type == "list_item":
            story.append(
                ListFlowable(
                    [Paragraph(block.text.replace("&", "&amp;"), body)],
                    bulletType="bullet",
                    leftIndent=16,
                )
            )
        elif block.type == "blockquote":
            story.append(Paragraph(block.text.replace("&", "&amp;"), quote))
        else:
            story.append(Paragraph(block.text.replace("&", "&amp;"), body))
        story.append(Spacer(1, 6))
    SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    ).build(story)
    return buffer.getvalue()


def export_document(blocks: list[Block], fmt: str) -> tuple[bytes, str]:
    """Render blocks to (bytes, filename). Raises ValidationError on unknown fmt."""
    if fmt not in MEDIA_TYPES:
        raise ValidationError(
        f"Unknown export format '{fmt}'. Allowed: docx, pdf.", "INVALID_FORMAT"
    )
    content = _to_docx(blocks) if fmt == "docx" else _to_pdf(blocks)
    return content, f"aiverse-export.{fmt}"
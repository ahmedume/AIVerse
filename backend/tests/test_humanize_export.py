# tests/test_humanize_export.py
# Purpose: humanizer + export service and API tests.

import io
import json

import pytest
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.blocks import Block
from app.core.exceptions import ValidationError
from app.services import export_service, humanize_service, parse_service


def _frames(text: str) -> list[dict]:
    return [
        json.loads(line.split("data: ", 1)[1])
        for line in text.strip().split("\n\n")
        if line.startswith("data: ")
    ]


def test_level_profiles():
    assert humanize_service.level_profile(1)[0] == humanize_service._PROFILES["human"]
    assert humanize_service.level_profile(2)[0] == humanize_service._PROFILES["human"]
    assert humanize_service.level_profile(3)[0] == humanize_service._PROFILES["balanced"]
    assert humanize_service.level_profile(5)[0] == humanize_service._PROFILES["balanced"]
    assert humanize_service.level_profile(6)[0] == humanize_service._PROFILES["corporate"]
    assert humanize_service.level_profile(7)[0] == humanize_service._PROFILES["corporate"]
    assert humanize_service.level_profile(1)[1] > humanize_service.level_profile(7)[1]


def test_prompt_contract_protects_facts():
    prompt = humanize_service._prompt_for("Revenue grew 12% in 2024.", 3)
    assert "12%" in prompt and "2024" in prompt


async def test_humanize_stream_events(monkeypatch):
    class FakeModel:
        async def astream(self, messages):
            yield type("Chunk", (), {"content": "Rewritten "})()
            yield type("Chunk", (), {"content": "text."})()

    monkeypatch.setattr(humanize_service, "get_model_chain", lambda *a, **k: [FakeModel()])
    blocks = parse_service.resolve_source(
        {"text": "# Keep Me\n\nOriginal paragraph here to rewrite."}
    )[1]
    events = []
    async for frame in humanize_service.humanize_stream(
        {"text": "# Keep Me\n\nOriginal paragraph here to rewrite."}, 3
    ):
        events.append(json.loads(frame.split("data: ", 1)[1]))
    assert events[0]["event"] == "meta"
    names = [e["event"] for e in events]
    assert names[1] == "block_start"
    assert "token" in names and "block_end" in names
    assert events[-1]["event"] == "done"
    assert events[-1]["data"]["level"] == 3
    rewritten = events[-1]["data"]["blocks"]
    assert rewritten[0]["text"] == "Rewritten text."
    assert "# Keep Me" not in rewritten[0]["text"]


async def test_humanize_stream_keeps_heading_verbatim(monkeypatch):
    class FakeModel:
        async def astream(self, messages):
            yield type("Chunk", (), {"content": "New body."})()

    monkeypatch.setattr(humanize_service, "get_model_chain", lambda *a, **k: [FakeModel()])
    events = []
    async for frame in humanize_service.humanize_stream({"text": "# The Title\n\nBody text."}, 1):
        events.append(json.loads(frame.split("data: ", 1)[1]))
    blocks = events[-1]["data"]["blocks"]
    assert len(blocks) == 1
    assert blocks[0]["type"] == "paragraph"


async def test_humanize_stream_falls_back_to_second_provider(monkeypatch):
    class BrokenModel:
        async def astream(self, messages):
            raise RuntimeError("boom")
            yield  # pragma: no cover

    class GoodModel:
        async def astream(self, messages):
            yield type("Chunk", (), {"content": "Recovered text."})()

    monkeypatch.setattr(humanize_service, "get_model_chain", lambda *a, **k: [BrokenModel(), GoodModel()])
    events = []
    async for frame in humanize_service.humanize_stream({"text": "Original words here."}, 4):
        events.append(json.loads(frame.split("data: ", 1)[1]))
    assert events[-1]["data"]["blocks"][0]["text"] == "Recovered text."


async def test_humanize_stream_all_models_fail_falls_back(monkeypatch):
    class BrokenModel:
        async def astream(self, messages):
            raise RuntimeError("boom")
            yield  # pragma: no cover

    monkeypatch.setattr(humanize_service, "get_model_chain", lambda *a, **k: [BrokenModel()])
    events = []
    async for frame in humanize_service.humanize_stream({"text": "Original words here."}, 4):
        events.append(json.loads(frame.split("data: ", 1)[1]))
    assert events[-1]["data"]["blocks"][0]["text"] == "Original words here."


def test_export_docx_preserves_structure():
    blocks = [
        Block(index=0, type="heading", text="Title", level=1),
        Block(index=1, type="paragraph", text="Some body."),
        Block(index=2, type="list_item", text="First point"),
        Block(index=3, type="blockquote", text="A quote"),
    ]
    content, filename = export_service.export_document(blocks, "docx")
    assert filename.endswith(".docx")
    assert content[:2] == b"PK"
    doc = DocxDocument(io.BytesIO(content))
    styles = [p.style.name for p in doc.paragraphs]
    assert styles[0].startswith("Heading")
    assert "List Bullet" in styles
    assert len(doc.paragraphs) == 4


def test_export_pdf_is_valid_and_has_text():
    blocks = [
        Block(index=0, type="heading", text="Report", level=1),
        Block(index=1, type="paragraph", text="The quick brown fox jumped over the fence."),
    ]
    content, filename = export_service.export_document(blocks, "pdf")
    assert filename.endswith(".pdf")
    assert content[:5] == b"%PDF-"
    text = "".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(content)).pages)
    assert "Report" in text


def test_export_invalid_format():
    with pytest.raises(ValidationError) as exc:
        export_service.export_document([], "exe")
    assert exc.value.code == "INVALID_FORMAT"


def test_humanize_api_streams(client, monkeypatch):
    class FakeModel:
        async def astream(self, messages):
            yield type("Chunk", (), {"content": "Smoother prose."})()

    monkeypatch.setattr(humanize_service, "get_model_chain", lambda *a, **k: [FakeModel()])
    resp = client.post("/api/humanize", json={"source": {"text": "Stiff original sentence here."}, "level": 5})
    assert resp.status_code == 200
    events = _frames(resp.text)
    assert events[-1]["event"] == "done"
    assert events[-1]["data"]["blocks"][0]["text"] == "Smoother prose."


def test_humanize_api_invalid_level(client):
    resp = client.post("/api/humanize", json={"source": {"text": "Hello."}, "level": 9})
    assert resp.status_code == 422


def test_export_api_docx_download(client):
    resp = client.post("/api/export", json={"source": {"text": "# Head\n\nBody paragraph."}, "format": "docx"})
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("application/vnd.openxmlformats")
    assert "attachment" in resp.headers["content-disposition"]
    assert resp.content[:2] == b"PK"


def test_export_api_invalid_format(client):
    resp = client.post("/api/export", json={"source": {"text": "Hello."}, "format": "exe"})
    assert resp.status_code == 422
    assert resp.json()["error"]["code"] == "INVALID_FORMAT"
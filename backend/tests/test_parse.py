# tests/test_parse.py
# Purpose: parse_service unit tests — formats, sniffing, storage roundtrip, errors.

import io

import pytest
from docx import Document

from app.core.exceptions import (
    EmptyDocumentError,
    FileTooLargeError,
    NotFoundError,
    ParseFailedError,
    UnsupportedFileTypeError,
)
from app.services import parse_service


def _types(blocks):
    return [b.type for b in blocks]


def test_txt_paragraphs_and_heading():
    raw = b"INTRODUCTION\n\nFirst paragraph here.\nSecond sentence.\n\nAnother paragraph."
    ext, blocks = parse_service.parse_document(raw, "doc.txt")
    assert ext == "txt"
    assert _types(blocks) == ["heading", "paragraph", "paragraph"]
    assert blocks[0].text == "INTRODUCTION"
    assert blocks[0].level == 1


def test_md_structure():
    raw = b"# Title\n\n## Sub\n\n- item one\n- item two\n\n> a quote\n\nBody text."
    ext, blocks = parse_service.parse_document(raw, "notes.md")
    assert ext == "md"
    assert _types(blocks) == ["heading", "heading", "list_item", "list_item", "blockquote", "paragraph"]
    assert blocks[0].level == 1
    assert blocks[1].level == 2


def test_json_list_of_strings():
    raw = b'["First para.", "Second para."]'
    ext, blocks = parse_service.parse_document(raw, "data.json")
    assert ext == "json"
    assert len(blocks) == 2
    assert blocks[0].text == "First para."


def test_pdf_roundtrip():
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    pdf = canvas.Canvas(buf)
    pdf.drawString(72, 720, "Hello from a PDF.")
    pdf.save()
    ext, blocks = parse_service.parse_document(buf.getvalue(), "paper.pdf")
    assert ext == "pdf"
    assert any("PDF" in b.text for b in blocks)


def test_docx_roundtrip():
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("Report Title", level=1)
    doc.add_paragraph("Some body text here.")
    doc.add_paragraph("Bullet content", style="List Bullet")
    doc.save(buf)
    ext, blocks = parse_service.parse_document(buf.getvalue(), "report.docx")
    assert ext == "docx"
    assert _types(blocks) == ["heading", "paragraph", "list_item"]
    assert blocks[0].level == 1


def test_unsupported_extension():
    with pytest.raises(UnsupportedFileTypeError):
        parse_service.parse_document(b"hello", "virus.exe")


def test_disguised_pdf_rejected():
    with pytest.raises(ParseFailedError):
        parse_service.parse_document(b"%PDF-1.7 garbage not really pdf", "innocent.txt")


def test_binary_text_rejected():
    with pytest.raises(ParseFailedError):
        parse_service.parse_document(bytes([0, 1, 2, 3, 4, 255]) * 100, "notes.txt")


def test_empty_document():
    with pytest.raises(EmptyDocumentError):
        parse_service.parse_document(b"   \n\n  ", "empty.txt")


def test_too_large():
    big = b"x" * (20 * 1024 * 1024 + 1)
    with pytest.raises(FileTooLargeError):
        parse_service.parse_document(big, "huge.txt")


def test_save_load_delete_roundtrip():
    file_id = parse_service.create_file_id()
    raw = b"Alpha.\n\nBeta."
    ext, blocks = parse_service.parse_document(raw, "a.txt")
    parse_service.save_document(file_id, "a.txt", ext, raw, blocks)
    loaded_name, loaded_blocks = parse_service.load_document(file_id)
    assert loaded_name == "a.txt"
    assert len(loaded_blocks) == 2
    assert file_id in {item["id"] for item in parse_service.list_documents()}
    parse_service.delete_document(file_id)
    assert file_id not in {item["id"] for item in parse_service.list_documents()}
    with pytest.raises(NotFoundError):
        parse_service.load_document(file_id)


def test_delete_missing_raises():
    with pytest.raises(NotFoundError):
        parse_service.delete_document("deadbeefdeadbeef")


def test_bad_file_id_rejected():
    with pytest.raises(NotFoundError):
        parse_service.load_document("../../etc/passwd")